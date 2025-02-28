import os
import requests
from docx import Document

def get_pmc_metadata(query, max_results=10, output_file="pmc_metadata.docx"):
    # 如果指定了目录路径，确保目录存在
    output_dir = os.path.dirname(output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    base_url = "https://www.ebi.ac.uk/europepmc/webservices/rest/search"
    params = {
        "query": query,
        "resultType": "core",
        "pageSize": max_results,
        "format": "json"
    }

    try:
        response = requests.get(base_url, params=params)
        response.raise_for_status()
        data = response.json()
    except requests.exceptions.RequestException as e:
        print(f"Error fetching data from Europe PMC: {e}")
        print("Please check the URL and your network connection. If the problem persists, try again later.")
        return

    if "resultList" not in data or "result" not in data["resultList"]:
        print("No results found or invalid response format.")
        return

    # 创建一个新的Word文档
    document = Document()
    document.add_heading(f"PubMed Central (PMC) Metadata for Query: {query}", level=1)

    for article in data["resultList"]["result"]:
        title = article.get("title", "No title")
        authors = article.get("authorList", {}).get("author", [])
        authors = ", ".join([author.get("fullName", "No name") for author in authors])
        journal = article.get("journalTitle", "No journal")
        abstract = article.get("abstractText", "No abstract")

        # 将文章信息添加到Word文档
        document.add_heading(title, level=2)
        document.add_paragraph(f"Authors: {authors}")
        document.add_paragraph(f"Journal: {journal}")
        document.add_paragraph(f"Abstract: {abstract}")
        document.add_paragraph("-" * 50)

    # 保存Word文档
    document.save(output_file)
    print(f"Metadata saved to {output_file}")

# 示例：指定保存路径
output_file = r"C:\Users\Administrator\Desktop\pmc_metadata.docx"
get_pmc_metadata("cancer", max_results=5, output_file=output_file)