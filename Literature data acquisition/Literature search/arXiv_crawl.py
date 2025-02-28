import os
import requests
import xml.etree.ElementTree as ET
from docx import Document

def get_arxiv_metadata(query, max_results=10, output_file="arxiv_metadata.docx"):
    # 如果指定了目录路径，确保目录存在
    output_dir = os.path.dirname(output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    base_url = "http://export.arxiv.org/api/query"  # 修复链接格式
    params = {
        "search_query": query,
        "max_results": max_results
    }

    try:
        response = requests.get(base_url, params=params)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(f"Error fetching data from arXiv: {e}")
        print("Please check the URL and your network connection. If the problem persists, try again later.")
        return

    try:
        # 解析XML数据
        root = ET.fromstring(response.content)
    except ET.ParseError as e:
        print(f"Error parsing XML response: {e}")
        print("Please check the URL and ensure the response is in valid XML format.")
        return

    # 创建一个新的Word文档
    document = Document()
    document.add_heading(f"arXiv Metadata for Query: {query}", level=1)

    for entry in root.findall("{http://www.w3.org/2005/Atom}entry"):
        title = entry.find("{http://www.w3.org/2005/Atom}title").text if entry.find("{http://www.w3.org/2005/Atom}title") is not None else "No title"
        authors = ", ".join([author.find("{http://www.w3.org/2005/Atom}name").text for author in entry.findall("{http://www.w3.org/2005/Atom}author")])
        summary = entry.find("{http://www.w3.org/2005/Atom}summary").text if entry.find("{http://www.w3.org/2005/Atom}summary") is not None else "No summary"
        published = entry.find("{http://www.w3.org/2005/Atom}published").text if entry.find("{http://www.w3.org/2005/Atom}published") is not None else "No date"

        # 将文章信息添加到Word文档
        document.add_heading(title, level=2)
        document.add_paragraph(f"Authors: {authors}")
        document.add_paragraph(f"Published: {published}")
        document.add_paragraph(f"Summary: {summary}")
        document.add_paragraph("-" * 50)

    # 保存Word文档
    document.save(output_file)
    print(f"Metadata saved to {output_file}")

# 示例：获取与“quantum computing”相关的文献
output_file = r"C:\Users\Administrator\Desktop\arxiv_metadata.docx"
get_arxiv_metadata("quantum computing", max_results=5, output_file=output_file)