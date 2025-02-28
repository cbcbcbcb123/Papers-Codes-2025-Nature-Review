import requests
from bs4 import BeautifulSoup
from docx import Document
import os

def get_google_scholar_metadata(query, max_results=10):
    base_url = "https://scholar.google.com/scholar"
    params = {
        "q": query,
        "num": max_results
    }

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3"
    }

    try:
        response = requests.get(base_url, params=params, headers=headers, timeout=10)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(f"Error fetching data from Google Scholar: {e}")
        return []

    soup = BeautifulSoup(response.text, 'html.parser')
    results = soup.find_all('div', class_='gs_r gs_or gs_scl')

    metadata_list = []
    for result in results:
        title = result.find('h3', class_='gs_rt').text.strip() if result.find('h3', class_='gs_rt') else "No title"
        link = result.find('a')['href'] if result.find('a') else "No link"
        authors = result.find('div', class_='gs_a').text.strip() if result.find('div', class_='gs_a') else "No authors"
        snippet = result.find('div', class_='gs_rs').text.strip() if result.find('div', class_='gs_rs') else "No snippet"

        metadata = {
            "title": title,
            "link": link,
            "authors": authors,
            "snippet": snippet
        }
        metadata_list.append(metadata)

    return metadata_list

def save_to_word(results, output_path):
    # 创建一个新的 Word 文档
    doc = Document()
    doc.add_heading("Google Scholar Search Results", level=1)
    doc.add_paragraph(f"Query: {query}")

    for result in results:
        doc.add_heading(result['title'], level=2)
        doc.add_paragraph(f"Authors: {result['authors']}")
        doc.add_paragraph(f"Link: {result['link']}")
        doc.add_paragraph(f"Snippet: {result['snippet']}")
        doc.add_paragraph("-" * 50)

    # 保存文档
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Results saved to {output_path}")

# 示例：获取与“机器学习”相关的文献并保存到 Word 文档
query = "机器学习"
results = get_google_scholar_metadata(query, max_results=5)

if results:
    output_path = r"C:\Users\Administrator\Desktop\google_scholar.docx"  # 指定保存路径
    save_to_word(results, output_path)
else:
    print("No results found.")