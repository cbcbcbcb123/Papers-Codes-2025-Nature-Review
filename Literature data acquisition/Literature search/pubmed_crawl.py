import os
import requests
import xml.etree.ElementTree as ET
from docx import Document  # 用于创建Word文档

def get_pubmed_metadata(query, max_results=10, output_file="pubmed_metadata.docx"):
    # 如果指定了目录路径，确保目录存在
    output_dir = os.path.dirname(output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
    params = {
        "db": "pubmed",
        "term": query,
        "retmax": max_results,
        "retmode": "xml"
    }
    try:
        response = requests.get(base_url, params=params)
        response.raise_for_status()
        root = ET.fromstring(response.content)
    except requests.exceptions.RequestException as e:
        print(f"Error fetching data from PubMed: {e}")
        return

    pmids = [id.text for id in root.findall(".//Id")]
    print(f"Found {len(pmids)} articles.")

    # 创建一个新的Word文档
    document = Document()
    document.add_heading(f"PubMed Metadata for Query: {query}", level=1)

    for pmid in pmids:
        fetch_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?db=pubmed&id={pmid}&retmode=xml"
        try:
            fetch_response = requests.get(fetch_url)
            fetch_response.raise_for_status()
            article_root = ET.fromstring(fetch_response.content)
        except requests.exceptions.RequestException as e:
            print(f"Error fetching article {pmid}: {e}")
            print(f"Failed URL: {fetch_url}")
            continue

        title = article_root.find(".//ArticleTitle").text if article_root.find(".//ArticleTitle") is not None else "Title not available"
        authors = ", ".join([
            f"{author.find('LastName').text if author.find('LastName') is not None else ''} {author.find('ForeName').text if author.find('ForeName') is not None else ''}".strip()
            for author in article_root.findall(".//Author")
        ])
        authors = authors if authors else "Authors not available"
        abstract = article_root.find(".//AbstractText").text if article_root.find(".//AbstractText") is not None else "Abstract not available"

        # 将文章信息添加到Word文档
        document.add_heading(title, level=2)
        document.add_paragraph(f"Authors: {authors}")
        document.add_paragraph(f"Abstract: {abstract}")
        document.add_paragraph("-" * 50)

    # 保存Word文档
    document.save(output_file)
    print(f"Metadata saved to {output_file}")

# 示例：指定保存路径
output_file = r"C:\Users\Administrator\Desktop\pubmed_metadata.docx"
get_pubmed_metadata("COVID-19", max_results=5, output_file=output_file)